"""
knowledge_base_router.py — Knowledge Base API Router

Handles:
  - Multi-format input (text, questionnaire, file upload: PDF/DOCX/TXT)
  - AI-powered extraction of structured company parameters
  - CRUD for structured knowledge base data
"""

import json
import os
import re
import tempfile
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import get_db
from models import CompanySettings

router = APIRouter(prefix="/api/v1/knowledge-base", tags=["knowledge-base"])


# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------

class StructuredKB(BaseModel):
    """The five structured parameters we extract from any input."""
    company_description: Optional[str] = ""
    product_offering: Optional[str] = ""
    target_customers: Optional[str] = ""
    value_proposition: Optional[str] = ""
    messaging_tone: Optional[str] = ""


class TextExtractRequest(BaseModel):
    """Send raw text / paragraph for AI extraction."""
    text: str
    openai_api_key: Optional[str] = None


class QuestionnaireRequest(BaseModel):
    """Pre-filled questionnaire answers."""
    company_description: Optional[str] = ""
    product_offering: Optional[str] = ""
    target_customers: Optional[str] = ""
    value_proposition: Optional[str] = ""
    messaging_tone: Optional[str] = ""
    openai_api_key: Optional[str] = None


class SaveStructuredRequest(BaseModel):
    """Save / update the structured KB directly."""
    company_description: Optional[str] = ""
    product_offering: Optional[str] = ""
    target_customers: Optional[str] = ""
    value_proposition: Optional[str] = ""
    messaging_tone: Optional[str] = ""


# ---------------------------------------------------------------------------
# File parsing helpers
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdfium2."""
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(file_bytes)
    text_parts = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        textpage = page.get_textpage()
        text_parts.append(textpage.get_text_range())
        textpage.close()
        page.close()
    pdf.close()
    return "\n".join(text_parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    import docx
    import io
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text files."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, Exception):
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# AI extraction
# ---------------------------------------------------------------------------

EXTRACTION_PROMPT = """You are a business analyst AI. Analyze the following company information and extract structured data.

INPUT TEXT:
\"\"\"
{text}
\"\"\"

Extract the following parameters from the text. If a parameter cannot be determined from the text, write "Not specified" for that field.

Return your response as a valid JSON object with EXACTLY these keys:
{{
  "company_description": "A clear 2-3 sentence description of what the company does",
  "product_offering": "The products or services the company sells",
  "target_customers": "The ideal customer profile / target audience",
  "value_proposition": "What makes the company different from competitors",
  "messaging_tone": "The communication style (e.g., formal, casual, friendly, authoritative, professional)"
}}

IMPORTANT: Return ONLY the JSON object. No other text before or after."""


def _call_openai_extract(text: str, api_key: str) -> dict:
    """Call OpenAI to extract structured company data from raw text."""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a precise business analyst. Always respond with valid JSON only."},
                {"role": "user", "content": EXTRACTION_PROMPT.format(text=text[:4000])},
            ],
            max_tokens=800,
            temperature=0.3,
        )
        raw = response.choices[0].message.content.strip()

        # Try to parse JSON from the response
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(raw)
    except json.JSONDecodeError:
        return {
            "company_description": text[:500],
            "product_offering": "Not specified",
            "target_customers": "Not specified",
            "value_proposition": "Not specified",
            "messaging_tone": "Not specified",
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI extraction failed: {str(e)}")


REFINE_PROMPT = """You are a business analyst AI. The user has answered a questionnaire about their company.
Refine and improve their answers to be clear, professional, and well-structured.

User's answers:
- Company Description: {company_description}
- Product Offering: {product_offering}
- Target Customers: {target_customers}
- Value Proposition: {value_proposition}
- Messaging Tone: {messaging_tone}

Return your response as a valid JSON object with EXACTLY these keys:
{{
  "company_description": "refined description",
  "product_offering": "refined product info",
  "target_customers": "refined target customer profile",
  "value_proposition": "refined value proposition",
  "messaging_tone": "refined tone description"
}}

Rules:
- Keep the user's intent but make each field clearer and more professional
- If a field is empty or says "Not specified", keep it as "Not specified"
- Return ONLY the JSON object. No other text."""


def _call_openai_refine(answers: dict, api_key: str) -> dict:
    """Call OpenAI to refine questionnaire answers."""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a precise business analyst. Always respond with valid JSON only."},
                {"role": "user", "content": REFINE_PROMPT.format(**answers)},
            ],
            max_tokens=800,
            temperature=0.3,
        )
        raw = response.choices[0].message.content.strip()
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(raw)
    except json.JSONDecodeError:
        return answers  # Return original if parsing fails
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI refinement failed: {str(e)}")


# ---------------------------------------------------------------------------
# Helper to persist structured data
# ---------------------------------------------------------------------------

def _save_structured_to_db(db: Session, data: dict) -> CompanySettings:
    """Persist structured KB fields to the database."""
    settings = db.query(CompanySettings).first()

    # Build a combined knowledge_base_text for backward compatibility
    combined_text = (
        f"Company Description: {data.get('company_description', '')}\n\n"
        f"Product Offering: {data.get('product_offering', '')}\n\n"
        f"Target Customers: {data.get('target_customers', '')}\n\n"
        f"Value Proposition: {data.get('value_proposition', '')}\n\n"
        f"Messaging Tone: {data.get('messaging_tone', '')}"
    )

    if not settings:
        settings = CompanySettings(
            knowledge_base_text=combined_text,
            company_description=data.get("company_description", ""),
            product_offering=data.get("product_offering", ""),
            target_customers=data.get("target_customers", ""),
            value_proposition=data.get("value_proposition", ""),
            messaging_tone=data.get("messaging_tone", ""),
        )
        db.add(settings)
    else:
        settings.knowledge_base_text = combined_text
        settings.company_description = data.get("company_description", "")
        settings.product_offering = data.get("product_offering", "")
        settings.target_customers = data.get("target_customers", "")
        settings.value_proposition = data.get("value_proposition", "")
        settings.messaging_tone = data.get("messaging_tone", "")

    db.commit()
    db.refresh(settings)
    return settings


# ---------------------------------------------------------------------------
# API Endpoints
# ---------------------------------------------------------------------------

@router.get("")
def get_knowledge_base(db: Session = Depends(get_db)):
    """Fetch the full knowledge base (structured + raw text)."""
    settings = db.query(CompanySettings).first()
    if not settings:
        return {
            "content": "",
            "structured": {
                "company_description": "",
                "product_offering": "",
                "target_customers": "",
                "value_proposition": "",
                "messaging_tone": "",
            },
        }
    return {
        "content": settings.knowledge_base_text or "",
        "structured": {
            "company_description": settings.company_description or "",
            "product_offering": settings.product_offering or "",
            "target_customers": settings.target_customers or "",
            "value_proposition": settings.value_proposition or "",
            "messaging_tone": settings.messaging_tone or "",
        },
    }


@router.post("")
def update_knowledge_base_legacy(data: dict, db: Session = Depends(get_db)):
    """Legacy endpoint — save plain text (backward compat)."""
    content = data.get("content", "")
    settings = db.query(CompanySettings).first()
    if not settings:
        settings = CompanySettings(knowledge_base_text=content)
        db.add(settings)
    else:
        settings.knowledge_base_text = content
    db.commit()
    db.refresh(settings)
    return {"status": "success", "content": settings.knowledge_base_text}


@router.post("/extract-text")
def extract_from_text(req: TextExtractRequest, db: Session = Depends(get_db)):
    """Extract structured data from raw text / paragraph using AI."""
    api_key = req.openai_api_key or os.getenv("OPENAI_API_KEY", "")
    if not api_key or api_key == "sk-your-key-here":
        raise HTTPException(status_code=400, detail="OpenAI API key is required. Set it in .env or provide it in the request.")

    structured = _call_openai_extract(req.text, api_key)
    return {"status": "success", "structured": structured}


@router.post("/extract-file")
async def extract_from_file(
    file: UploadFile = File(...),
    openai_api_key: str = Form(""),
    db: Session = Depends(get_db),
):
    """Upload a PDF / DOCX / TXT file, extract text, then use AI to structure it."""
    api_key = openai_api_key or os.getenv("OPENAI_API_KEY", "")
    if not api_key or api_key == "sk-your-key-here":
        raise HTTPException(status_code=400, detail="OpenAI API key is required.")

    filename = (file.filename or "").lower()
    file_bytes = await file.read()

    if filename.endswith(".pdf"):
        raw_text = _extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        raw_text = _extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        raw_text = _extract_text_from_txt(file_bytes)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Accepted: .pdf, .docx, .txt")

    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the uploaded file.")

    structured = _call_openai_extract(raw_text, api_key)
    return {"status": "success", "extracted_text": raw_text[:2000], "structured": structured}


@router.post("/refine-questionnaire")
def refine_questionnaire(req: QuestionnaireRequest, db: Session = Depends(get_db)):
    """Take questionnaire answers and refine them using AI."""
    api_key = req.openai_api_key or os.getenv("OPENAI_API_KEY", "")
    if not api_key or api_key == "sk-your-key-here":
        raise HTTPException(status_code=400, detail="OpenAI API key is required.")

    answers = {
        "company_description": req.company_description or "Not specified",
        "product_offering": req.product_offering or "Not specified",
        "target_customers": req.target_customers or "Not specified",
        "value_proposition": req.value_proposition or "Not specified",
        "messaging_tone": req.messaging_tone or "Not specified",
    }

    refined = _call_openai_refine(answers, api_key)
    return {"status": "success", "structured": refined}


@router.post("/save-structured")
def save_structured(req: SaveStructuredRequest, db: Session = Depends(get_db)):
    """Save the final structured KB to the database."""
    data = req.dict()
    settings = _save_structured_to_db(db, data)
    return {
        "status": "success",
        "structured": {
            "company_description": settings.company_description or "",
            "product_offering": settings.product_offering or "",
            "target_customers": settings.target_customers or "",
            "value_proposition": settings.value_proposition or "",
            "messaging_tone": settings.messaging_tone or "",
        },
    }
